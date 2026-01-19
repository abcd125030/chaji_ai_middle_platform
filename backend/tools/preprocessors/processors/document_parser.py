from tools.core.base import BaseTool
from tools.core.exceptions import ToolExecutionError
from typing import Dict, Any
from docx import Document
import os,logging
from ..parsers.pdf_parser_internal import PDFParserInternalTool
logger = logging.getLogger(__name__)

class DocumentParserTool(BaseTool):
    """
    文档解析工具：支持将.docx和.pdf文件内容转换为Markdown文本。
    """
    
    def __init__(self, config: Dict[str, Any] = None):
        super().__init__(config)
        self.pdf_parser = PDFParserInternalTool(config)
    def get_input_schema(self) -> Dict[str, Any]:
        return {
            "type": "object",
            "properties": {
                "file_path": {
                    "type": "string",
                    "description": "要解析的文档文件路径（支持.docx和.pdf格式）"
                }
            },
            "required": ["file_path"]
        }

    def execute(self, tool_input: Dict[str, Any]) -> Dict[str, Any]:
        file_path = tool_input.get('file_path')
        
        if not file_path:
            return {"status": "error", "message": "文件路径未提供。"}
        
        if not os.path.exists(file_path):
            return {"status": "error", "message": f"文件不存在: {file_path}"}
        
        # 检查文件格式
        file_ext = os.path.splitext(file_path.lower())[1]
        
        if file_ext == '.pdf':
            # 使用PDF解析器
            result = self.pdf_parser.execute(tool_input)
            # 为PDF结果也生成摘要
            if result.get('status') == 'success':
                content = result.get('markdown_content', result.get('content', ''))
                summary = self._generate_summary(content)
                result['summary'] = summary
            return result
        elif file_ext == '.docx':
            # 使用原有的docx解析逻辑
            pass  # 继续执行下面的docx处理
        else:
            return {"status": "error", "message": f"不支持的文件格式，只支持.docx和.pdf: {file_path}"}
 
        try:
            doc = Document(file_path)
            markdown_content = self._convert_docx_to_markdown(doc)
            
            # 生成文档摘要
            summary = self._generate_summary(markdown_content)
            
            return {
                "status": "success",
                "message": f"文件 '{os.path.basename(file_path)}' 已成功转换为Markdown。",
                "file_path": file_path,
                "markdown_content": markdown_content,
                "summary": summary
            }
 
        except Exception as e:
            return {"status": "error", "message": f"文档解析失败: {str(e)}", "file_path": file_path}

    def _convert_docx_to_markdown(self, doc: Document) -> str:
        """
        将docx文档内容转换为Markdown格式。
        支持标题、段落、图片和表格的转换。
        """
        # 定义命名空间
        namespaces = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
            'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
            'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
        }
        
        markdown_lines = []
        
        for paragraph in doc.paragraphs:
            if not paragraph.text.strip() and not any(
                run._element.find('.//w:drawing', namespaces) is not None
                for run in paragraph.runs
            ):
                continue
                
            # 处理段落样式
            style_name = paragraph.style.name if paragraph.style else ''
            if style_name.startswith('Heading'):
                level = style_name.replace('Heading', '')
                try:
                    level = int(level)
                    markdown_lines.append('#' * level + ' ' + paragraph.text)
                    continue
                except ValueError:
                    pass
            
            # 处理图片
            has_image = False
            text = paragraph.text
            
            for run in paragraph.runs:
                drawing = run._element.find('.//w:drawing', namespaces)
                if drawing is not None:
                    has_image = True
                    # 图片处理占位符（原方法不处理文件路径）
                    markdown_lines.append('![image](图片占位符)')
            
            if not has_image and text.strip():
                # 处理文本样式
                if style_name == 'List Paragraph':
                    markdown_lines.append(f'- {text}')
                else:
                    markdown_lines.append(text)
        
        # 处理表格
        for table in doc.tables:
            table_md = []
            # 处理表头
            header_row = []
            for cell in table.rows[0].cells:
                header_row.append(cell.text.strip() or ' ')
            table_md.append('| ' + ' | '.join(header_row) + ' |')
            
            # 添加分隔行
            table_md.append('| ' + ' | '.join(['---'] * len(header_row)) + ' |')
            
            # 处理数据行
            for row in table.rows[1:]:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip() or ' ')
                table_md.append('| ' + ' | '.join(row_data) + ' |')
            
            # 在表格前后添加空行
            markdown_lines.append('')  # 表格前的空行
            markdown_lines.extend(table_md)
            markdown_lines.append('')  # 表格后的空行
        
        return '\n\n'.join(markdown_lines)
    
    def _generate_summary(self, content: str) -> str:
        """
        使用LLM生成文档摘要
        """
        try:
            from llm.llm_service import LLMService
            from django.conf import settings
            
            # 如果内容过长，截取前3000字符用于生成摘要
            max_content_length = 3000
            truncated_content = content[:max_content_length] if len(content) > max_content_length else content
            
            # 构建生成摘要的消息
            messages = [
                {
                    "role": "system",
                    "content": "你是一个专业的文档摘要助手。请为用户提供的文档内容生成一个简洁、准确的摘要。摘要应该包含文档的主要内容和关键信息，长度控制在200字以内。"
                },
                {
                    "role": "user", 
                    "content": f"请为以下文档生成摘要：\n\n{truncated_content}"
                }
            ]
            
            # 使用环境变量中的默认模型
            default_model = getattr(settings, 'DEFAULT_MODEL', 'qwen3-coder-plus')
            
            # 调用LLM服务生成摘要
            llm_service = LLMService()
            response = llm_service.internal.call_llm(
                model_name=default_model,
                messages=messages,
                temperature=0.3,
                max_tokens=400  # 确保能生成200字的中文摘要
            )
            
            # 提取摘要内容
            if response and isinstance(response, dict):
                if 'choices' in response and response['choices']:
                    summary = response['choices'][0].get('message', {}).get('content', '')
                    if summary:
                        return summary.strip()
            
            # 如果生成失败，返回内容的前200字作为简单摘要
            return content[:200] + "..." if len(content) > 200 else content
            
        except Exception as e:
            logger.warning(f"生成摘要失败: {str(e)}")
            # 生成摘要失败时，返回内容的前200字作为备用摘要
            return content[:200] + "..." if len(content) > 200 else content