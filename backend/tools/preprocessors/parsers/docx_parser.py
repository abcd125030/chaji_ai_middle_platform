"""
DOCX文档解析工具
================

将Microsoft Word文档（.docx）内容转换为结构化的Markdown格式。
支持提取文本、表格、列表、标题层级、图片信息等。
"""

import os
import io
import base64
import logging
from typing import Dict, Any, List, Optional
from tools.core.base import BaseTool
from tools.core.exceptions import ToolExecutionError

logger = logging.getLogger(__name__)


class DOCXParserTool(BaseTool):
    """
    DOCX文档解析工具：将Word文档内容转换为Markdown格式。
    
    功能特性：
    1. 提取文档文本内容，保留段落结构
    2. 识别并转换标题层级（Heading 1-6）
    3. 提取并格式化表格为Markdown表格
    4. 处理有序和无序列表
    5. 提取图片信息和说明文字
    6. 保留文本格式（加粗、斜体、下划线等）
    7. 提取文档元数据（作者、创建时间、修改时间等）
    """
    
    def get_input_schema(self) -> Dict[str, Any]:
        return {
            "type": "object",
            "properties": {
                "file_path": {
                    "type": "string",
                    "description": "要解析的DOCX文件路径"
                },
                "extract_images": {
                    "type": "boolean",
                    "description": "是否提取图片（base64编码）",
                    "default": False
                },
                "include_metadata": {
                    "type": "boolean",
                    "description": "是否包含文档元数据",
                    "default": True
                },
                "preserve_formatting": {
                    "type": "boolean",
                    "description": "是否保留文本格式（加粗、斜体等）",
                    "default": True
                }
            },
            "required": ["file_path"]
        }

    def execute(self, tool_input: Dict[str, Any]) -> Dict[str, Any]:
        """执行DOCX文档解析"""
        file_path = tool_input.get('file_path')
        extract_images = tool_input.get('extract_images', False)
        include_metadata = tool_input.get('include_metadata', True)
        preserve_formatting = tool_input.get('preserve_formatting', True)
        
        # 验证输入
        if not file_path:
            return {"status": "error", "message": "文件路径未提供"}
        
        if not os.path.exists(file_path):
            return {"status": "error", "message": f"文件不存在: {file_path}"}
        
        if not file_path.lower().endswith(('.docx', '.doc')):
            return {"status": "error", "message": f"不支持的文件格式，只支持.docx/.doc: {file_path}"}
        
        try:
            # 解析文档
            result = self._parse_docx(
                file_path, 
                extract_images=extract_images,
                include_metadata=include_metadata,
                preserve_formatting=preserve_formatting
            )
            
            return {
                "status": "success",
                "message": f"文件 '{os.path.basename(file_path)}' 已成功解析",
                "file_path": file_path,
                **result
            }
            
        except ImportError as e:
            logger.error(f"缺少必要的依赖库: {str(e)}")
            return {
                "status": "error",
                "message": "需要安装python-docx库: pip install python-docx",
                "file_path": file_path
            }
        except Exception as e:
            logger.error(f"DOCX解析失败: {str(e)}", exc_info=True)
            return {
                "status": "error",
                "message": f"DOCX解析失败: {str(e)}",
                "file_path": file_path
            }
    
    def _parse_docx(self, file_path: str, extract_images: bool = False,
                    include_metadata: bool = True, 
                    preserve_formatting: bool = True) -> Dict[str, Any]:
        """解析DOCX文档的核心方法"""
        try:
            from docx import Document
            from docx.shared import Inches
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
            from docx.enum.style import WD_STYLE_TYPE
        except ImportError:
            raise ImportError("需要安装python-docx库: pip install python-docx")
        
        # 打开文档
        doc = Document(file_path)
        
        # 提取内容
        markdown_content = []
        images_data = []
        tables_data = []
        statistics = {
            "paragraphs": 0,
            "words": 0,
            "characters": 0,
            "tables": 0,
            "images": 0,
            "lists": 0,
            "headings": 0
        }
        
        # 处理文档元数据
        metadata = {}
        if include_metadata:
            metadata = self._extract_metadata(doc)
        
        # 遍历文档元素
        for element in doc.element.body:
            # 判断元素类型并处理
            if element.tag.endswith('p'):  # 段落
                para = self._element_to_paragraph(element, doc)
                if para:
                    content = self._process_paragraph(para, preserve_formatting)
                    if content:
                        markdown_content.append(content)
                        statistics["paragraphs"] += 1
                        
                        # 统计标题
                        if para.style.name.startswith('Heading'):
                            statistics["headings"] += 1
                        
                        # 统计单词
                        statistics["words"] += len(para.text.split())
                        statistics["characters"] += len(para.text)
                        
            elif element.tag.endswith('tbl'):  # 表格
                table = self._element_to_table(element, doc)
                if table:
                    table_md = self._process_table(table)
                    if table_md:
                        markdown_content.append(table_md)
                        tables_data.append(self._extract_table_data(table))
                        statistics["tables"] += 1
        
        # 处理图片
        if extract_images:
            images_data = self._extract_images(doc)
            statistics["images"] = len(images_data)
            
            # 在Markdown中添加图片引用
            for idx, img_data in enumerate(images_data):
                markdown_content.append(f"\n![图片{idx + 1}]({img_data.get('filename', f'image_{idx + 1}')})\n")
        
        # 组装最终内容
        final_content = "\n\n".join(markdown_content)
        
        result = {
            "content": final_content,
            "statistics": statistics,
            "tables": tables_data
        }
        
        if include_metadata:
            result["metadata"] = metadata
        
        if extract_images:
            result["images"] = images_data
        
        return result
    
    def _process_paragraph(self, paragraph, preserve_formatting: bool = True) -> str:
        """处理段落，转换为Markdown格式"""
        if not paragraph.text.strip():
            return ""
        
        # 检查是否是标题
        style_name = paragraph.style.name
        if style_name.startswith('Heading'):
            level = self._get_heading_level(style_name)
            return f"{'#' * level} {paragraph.text.strip()}"
        
        # 检查是否是列表项
        if style_name.startswith('List'):
            # 判断是有序还是无序列表
            if 'Number' in style_name:
                return f"1. {paragraph.text.strip()}"
            else:
                return f"- {paragraph.text.strip()}"
        
        # 处理普通段落
        if preserve_formatting:
            return self._apply_text_formatting(paragraph)
        else:
            return paragraph.text.strip()
    
    def _apply_text_formatting(self, paragraph) -> str:
        """应用文本格式（加粗、斜体等）"""
        formatted_text = []
        
        for run in paragraph.runs:
            text = run.text
            if not text:
                continue
            
            # 加粗
            if run.bold:
                text = f"**{text}**"
            
            # 斜体
            if run.italic:
                text = f"*{text}*"
            
            # 下划线（Markdown不直接支持，使用HTML标签）
            if run.underline:
                text = f"<u>{text}</u>"
            
            formatted_text.append(text)
        
        return "".join(formatted_text)
    
    def _process_table(self, table) -> str:
        """处理表格，转换为Markdown表格格式"""
        if not table.rows:
            return ""
        
        markdown_table = []
        
        # 处理表格行
        for row_idx, row in enumerate(table.rows):
            row_data = []
            for cell in row.cells:
                # 获取单元格文本，处理换行
                cell_text = cell.text.strip().replace('\n', ' ')
                row_data.append(cell_text)
            
            # 添加行
            markdown_table.append("| " + " | ".join(row_data) + " |")
            
            # 在第一行后添加分隔符
            if row_idx == 0:
                separator = "|"
                for _ in row.cells:
                    separator += " --- |"
                markdown_table.append(separator)
        
        return "\n".join(markdown_table)
    
    def _extract_table_data(self, table) -> Dict[str, Any]:
        """提取表格数据为结构化格式"""
        headers = []
        rows = []
        
        for row_idx, row in enumerate(table.rows):
            row_data = [cell.text.strip() for cell in row.cells]
            
            if row_idx == 0:
                headers = row_data
            else:
                rows.append(row_data)
        
        return {
            "headers": headers,
            "rows": rows,
            "row_count": len(table.rows),
            "column_count": len(table.columns) if hasattr(table, 'columns') else len(headers)
        }
    
    def _extract_images(self, doc) -> List[Dict[str, Any]]:
        """提取文档中的图片"""
        images = []
        
        # 获取文档中的所有图片关系
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                try:
                    # 获取图片数据
                    image_data = rel.target_part.blob
                    
                    # 获取图片扩展名
                    ext = rel.target_part.partname.split('.')[-1]
                    
                    # Base64编码
                    image_base64 = base64.b64encode(image_data).decode('utf-8')
                    
                    images.append({
                        "filename": f"image_{len(images) + 1}.{ext}",
                        "extension": ext,
                        "size": len(image_data),
                        "base64": image_base64
                    })
                except Exception as e:
                    logger.warning(f"提取图片失败: {str(e)}")
        
        return images
    
    def _extract_metadata(self, doc) -> Dict[str, Any]:
        """提取文档元数据"""
        metadata = {}
        
        try:
            core_props = doc.core_properties
            
            # 基本属性
            if core_props.author:
                metadata["author"] = core_props.author
            if core_props.title:
                metadata["title"] = core_props.title
            if core_props.subject:
                metadata["subject"] = core_props.subject
            if core_props.keywords:
                metadata["keywords"] = core_props.keywords
            if core_props.comments:
                metadata["comments"] = core_props.comments
            
            # 时间属性
            if core_props.created:
                metadata["created"] = core_props.created.isoformat()
            if core_props.modified:
                metadata["modified"] = core_props.modified.isoformat()
            if core_props.last_modified_by:
                metadata["last_modified_by"] = core_props.last_modified_by
            
            # 其他属性
            if core_props.revision:
                metadata["revision"] = core_props.revision
            if core_props.category:
                metadata["category"] = core_props.category
            
        except Exception as e:
            logger.warning(f"提取元数据失败: {str(e)}")
        
        return metadata
    
    def _get_heading_level(self, style_name: str) -> int:
        """根据样式名称获取标题级别"""
        # Heading 1 -> 1, Heading 2 -> 2, etc.
        try:
            level = int(style_name.split()[-1])
            return min(max(level, 1), 6)  # 确保在1-6范围内
        except:
            return 1
    
    def _element_to_paragraph(self, element, doc):
        """将XML元素转换为段落对象"""
        for paragraph in doc.paragraphs:
            if paragraph._element == element:
                return paragraph
        return None
    
    def _element_to_table(self, element, doc):
        """将XML元素转换为表格对象"""
        for table in doc.tables:
            if table._element == element:
                return table
        return None


# 用于直接测试的辅助函数
def parse_docx_file(file_path: str, output_format: str = "markdown") -> str:
    """
    便捷函数：解析DOCX文件并返回内容
    
    Args:
        file_path: DOCX文件路径
        output_format: 输出格式（markdown/text/json）
    
    Returns:
        解析后的内容
    """
    parser = DOCXParserTool()
    result = parser.execute({"file_path": file_path})
    
    if result["status"] == "success":
        if output_format == "markdown":
            return result.get("content", "")
        elif output_format == "json":
            import json
            return json.dumps(result, ensure_ascii=False, indent=2)
        else:
            # 纯文本，去除Markdown格式
            content = result.get("content", "")
            # 简单处理：去除#号和*号
            content = content.replace('#', '').replace('*', '')
            return content
    else:
        raise Exception(result["message"])


if __name__ == "__main__":
    # 命令行测试
    import sys
    
    if len(sys.argv) < 2:
        print("用法: python docx_parser.py <docx文件路径> [输出格式:markdown/text/json]")
        sys.exit(1)
    
    file_path = sys.argv[1]
    output_format = sys.argv[2] if len(sys.argv) > 2 else "markdown"
    
    try:
        content = parse_docx_file(file_path, output_format)
        print(content)
    except Exception as e:
        print(f"错误: {e}", file=sys.stderr)
        sys.exit(1)